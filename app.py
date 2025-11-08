"""
Streamlit App for Internal Knowledge Navigator
"""
import streamlit as st
import json
import os
import platform
import subprocess
import smtplib
import ssl
import socket
import re
from difflib import SequenceMatcher
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formataddr
from pathlib import Path
from typing import List, Dict, Optional, Tuple

from dotenv import load_dotenv
try:
    import tkinter as tk
    from tkinter import filedialog
except Exception:
    tk = None

from server.vector_db import VectorDatabaseManager
from server.rag_engine import RAGEngine
from pypdf import PdfReader
from docx import Document


load_dotenv()


SUPPORTED_SUFFIXES = [".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".gif"]
IGNORED_DIRECTORIES = {".venv", "__pycache__", ".idea", "node_modules", "venv"}


def parse_query_intent(query: str) -> Dict[str, any]:
    """
    Parse user query to determine intent and extract parameters.
    
    Returns:
        Dictionary with:
        - intent: 'email', 'search', 'summarize', 'query'
        - email: extracted email address (if email intent)
        - document_name: extracted document name (if search/email intent)
        - original_query: original query string
    """
    query_lower = query.lower().strip()
    result = {
        "intent": "query",  # default
        "email": None,
        "document_name": None,
        "original_query": query
    }
    
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Email intent detection - check for email keywords OR standalone "send" followed by potential filename
    email_keywords = ["send mail", "send email", "email", "mail", "send to", "send a mail", "send an email"]
    has_email_keyword = any(keyword in query_lower for keyword in email_keywords)
    # Also detect if query starts with "send" (likely email intent)
    starts_with_send = query_lower.strip().startswith("send ")
    
    if has_email_keyword or starts_with_send:
        result["intent"] = "email"
        # Extract email address using regex
        emails = re.findall(email_pattern, query)
        if emails:
            result["email"] = emails[0]
        
        # Try to extract document name from email query (e.g., "send contract.pdf to user@email.com")
        # Remove email keywords and email address
        doc_name = query
        # Remove email keywords first (multi-word phrases)
        for keyword in email_keywords:
            doc_name = re.sub(rf'\b{re.escape(keyword)}\b', '', doc_name, flags=re.IGNORECASE)
        # Remove standalone "send" keyword (must be whole word)
        doc_name = re.sub(r'\bsend\b', '', doc_name, flags=re.IGNORECASE)
        # Remove email address
        doc_name = re.sub(email_pattern, '', doc_name)
        # Remove common words
        doc_name = re.sub(r'\b(to|with|attached|attachment|file|document)\b', '', doc_name, flags=re.IGNORECASE)
        doc_name = doc_name.strip(' ,.')
        if doc_name:
            result["document_name"] = doc_name
    
    # Search intent detection
    search_keywords = ["search", "find", "look for", "locate", "where is", "show me"]
    if any(keyword in query_lower for keyword in search_keywords):
        result["intent"] = "search"
        # Try to extract document name (remove search keywords and clean up)
        doc_name = query
        for keyword in search_keywords:
            doc_name = re.sub(rf'\b{keyword}\b', '', doc_name, flags=re.IGNORECASE)
        doc_name = doc_name.strip()
        # Remove email if present
        doc_name = re.sub(email_pattern, '', doc_name)
        doc_name = doc_name.strip(' ,.')
        if doc_name:
            result["document_name"] = doc_name
    
    # Summarize intent detection
    summarize_keywords = ["summarize", "summary", "summarise", "brief", "overview"]
    if any(keyword in query_lower for keyword in summarize_keywords):
        result["intent"] = "summarize"
        # Try to extract document name from summarize query (e.g., "summarize contract.pdf")
        doc_name = query
        for keyword in summarize_keywords:
            doc_name = re.sub(rf'\b{keyword}\b', '', doc_name, flags=re.IGNORECASE)
        # Remove common words
        doc_name = re.sub(r'\b(all|the|documents|document|files|file)\b', '', doc_name, flags=re.IGNORECASE)
        doc_name = doc_name.strip(' ,.')
        if doc_name:
            result["document_name"] = doc_name
    
    return result


def get_env_int(var_name: str, default: int) -> int:
    """Safely parse an integer from an environment variable."""
    raw_value = os.getenv(var_name)
    if raw_value is None or not raw_value.strip():
        return default
    try:
        return int(raw_value.strip())
    except ValueError:
        st.warning(
            f"Invalid integer for environment variable '{var_name}'; using default {default}."
        )
        return default


def get_env_bool(var_name: str, default: bool) -> bool:
    """Safely parse a boolean from an environment variable."""
    raw_value = os.getenv(var_name)
    if raw_value is None or not raw_value.strip():
        return default
    return raw_value.strip().lower() in {"1", "true", "yes", "on"}


def browse_for_folder() -> str:
    """Open a native folder selection dialog and return the chosen path (or empty string)."""
    if tk is None:
        return ""
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        path = filedialog.askdirectory()
        root.destroy()
        return path or ""
    except Exception:
        return ""


def extract_text(file_path: Path) -> str:
    """Extract text from PDF/DOCX/TXT"""
    try:
        if file_path.suffix.lower() == ".pdf":
            reader = PdfReader(file_path)
            return "\n".join(page.extract_text() for page in reader.pages)
        elif file_path.suffix.lower() == ".docx":
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        elif file_path.suffix.lower() == ".txt":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        else:
            return ""
    except Exception as e:
        st.warning(f"Error reading {file_path.name}: {str(e)}")
        return ""


def should_ignore_path(path: Path) -> bool:
    """Check whether the path should be ignored (e.g., virtual env, cache)."""
    return any(part.lower() in IGNORED_DIRECTORIES for part in path.parts)


def list_supported_files(folder_path: str) -> List[Path]:
    """Return a sorted list of supported files within the folder."""
    files: List[Path] = []

    if not folder_path or not folder_path.strip():
        return files

    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        return files

    for file_path in folder.rglob("*"):
        if should_ignore_path(file_path):
            continue
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_SUFFIXES:
            files.append(file_path)

    return sorted(files)


def format_file_display(base_path: Path, file_path: Path) -> str:
    """Return a human-friendly display string for a file path."""
    try:
        relative = file_path.relative_to(base_path)
        display_path = str(relative)
    except ValueError:
        display_path = str(file_path)

    try:
        size_kb = file_path.stat().st_size / 1024
        size_part = f"{size_kb:.1f} KB"
    except OSError:
        size_part = "Unknown size"

    return f"{display_path} ({size_part})"


def get_folder_structure(folder_path: str, max_depth: int = 3, current_depth: int = 0) -> list:
    """Get folder structure as a list of strings"""
    structure = []
    if not folder_path or not folder_path.strip():
        return structure
    
    try:
        folder = Path(folder_path)
        if not folder.exists():
            return [f"❌ Path does not exist: {folder_path}"]
        
        if not folder.is_dir():
            return [f"❌ Not a directory: {folder_path}"]
        
        if current_depth >= max_depth:
            return structure
        
        # Get all items in the folder
        items = sorted(folder.iterdir(), key=lambda x: (x.is_file(), x.name.lower()))
        
        for item in items:
            if should_ignore_path(item):
                continue
            indent = "  " * current_depth
            if item.is_dir():
                structure.append(f"{indent}📁 {item.name}/")
                # Recursively get subdirectory structure
                sub_structure = get_folder_structure(str(item), max_depth, current_depth + 1)
                structure.extend(sub_structure)
            else:
                suffix = item.suffix.lower()
                if suffix in SUPPORTED_SUFFIXES:
                    structure.append(f"{indent}📄 {item.name} ✅")
                else:
                    structure.append(f"{indent}📄 {item.name}")
        
        return structure
    except PermissionError:
        return [f"❌ Permission denied: {folder_path}"]
    except Exception as e:
        return [f"❌ Error reading folder: {str(e)}"]


def get_supported_files_count(folder_path: str) -> dict:
    """Count supported files in the folder"""
    counts = {"pdf": 0, "docx": 0, "txt": 0, "total": 0}
    
    if not folder_path or not folder_path.strip():
        return counts
    
    try:
        folder = Path(folder_path)
        if not folder.exists() or not folder.is_dir():
            return counts
        
        for file_path in folder.rglob("*"):
            if should_ignore_path(file_path):
                continue
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                if suffix == ".pdf":
                    counts["pdf"] += 1
                    counts["total"] += 1
                elif suffix == ".docx":
                    counts["docx"] += 1
                    counts["total"] += 1
                elif suffix == ".txt":
                    counts["txt"] += 1
                    counts["total"] += 1
        
        return counts
    except Exception:
        return counts


def find_documents(folder_path: str, search_query: str) -> tuple[list, list]:
    """
    Search for documents by name in the folder
    
    Args:
        folder_path: Path to search in
        search_query: Document name to search for (case-insensitive, partial match)
    
    Returns:
        Tuple of (matching file paths, suggested similar file paths)
    """
    exact_results: list[Path] = []
    suggestions: list[tuple[Path, float]] = []

    if not folder_path or not search_query:
        return exact_results, []
    
    try:
        folder = Path(folder_path)
        if not folder.exists() or not folder.is_dir():
            return exact_results, []
        
        search_lower = search_query.lower().strip()
        all_supported: list[Path] = []

        # Search for files matching the query
        for file_path in folder.rglob("*"):
            if should_ignore_path(file_path):
                continue
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                if suffix in SUPPORTED_SUFFIXES:
                    all_supported.append(file_path)
                # Check if filename contains the search query
                if suffix in SUPPORTED_SUFFIXES and search_lower in file_path.name.lower():
                    exact_results.append(file_path)

        exact_results = sorted(exact_results)

        if exact_results:
            return exact_results, []

        # Build similarity suggestions if no direct matches
        if all_supported:
            for file_path in all_supported:
                similarity = SequenceMatcher(None, search_lower, file_path.name.lower()).ratio()
                if similarity >= 0.4:  # Only keep reasonably similar names
                    suggestions.append((file_path, similarity))

        suggestions = sorted(suggestions, key=lambda item: item[1], reverse=True)[:5]

        return [], suggestions
    except Exception as e:
        st.error(f"Error searching for documents: {str(e)}")
        return [], []


def open_file_location(file_path: Path):
    """
    Open file location in OS file explorer with file selected
    
    Args:
        file_path: Path to the file
    """
    try:
        system = platform.system()
        file_path_str = str(file_path.resolve())
        
        if system == "Windows":
            # Windows: explorer /select,"file_path"
            subprocess.run(['explorer', '/select,', file_path_str], check=False)
        elif system == "Darwin":  # macOS
            # macOS: open -R "file_path"
            subprocess.run(['open', '-R', file_path_str], check=False)
        elif system == "Linux":
            # Linux: try different file managers
            parent_dir = str(file_path.parent)
            # Try nautilus (GNOME)
            try:
                subprocess.run(['nautilus', '--select', file_path_str], check=False)
            except FileNotFoundError:
                # Try dolphin (KDE)
                try:
                    subprocess.run(['dolphin', '--select', file_path_str], check=False)
                except FileNotFoundError:
                    # Try thunar (XFCE) or just open parent directory
                    try:
                        subprocess.run(['thunar', parent_dir], check=False)
                    except FileNotFoundError:
                        # Fallback: open parent directory with xdg-open
                        subprocess.run(['xdg-open', parent_dir], check=False)
        else:
            st.warning(f"Unsupported OS: {system}. Cannot open file location.")
    except Exception as e:
        st.error(f"Error opening file location: {str(e)}")


def index_folder(folder_path: str, db: VectorDatabaseManager, force_reindex: bool = False) -> int:
    """
    Index all documents in the folder.
    Uses session state to track if folder is already indexed to avoid re-indexing.
    
    Args:
        folder_path: Path to folder to index
        db: VectorDatabaseManager instance
        force_reindex: If True, re-index even if already indexed
    
    Returns:
        Number of documents indexed
    """
    # Check if already indexed (using session state)
    index_key = f"indexed_folder_{folder_path}"
    if not force_reindex and index_key in st.session_state:
        if st.session_state[index_key] == folder_path:
            # Already indexed, just return count without showing message
            all_doc_ids = db.get_all_document_ids()
            return len(all_doc_ids) if all_doc_ids else 0
    
    # Show indexing message only when actually indexing
    with st.spinner("📊 Indexing documents..."):
        documents = {}
        count = 0
        
        folder = Path(folder_path)
        if not folder.exists():
            st.error(f"Folder not found: {folder_path}")
            return 0
        
        for file_path in folder.rglob("*"):
            if should_ignore_path(file_path):
                continue
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_SUFFIXES:
                content = extract_text(file_path)
                if content.strip():
                    doc_id = str(file_path).replace("\\", "_").replace("/", "_")
                    documents[doc_id] = {
                        "content": content,
                        "metadata": {
                            "filename": file_path.name,
                            "source": str(file_path),
                            "type": file_path.suffix[1:].lower()
                        }
                    }
                    count += 1
        
        if documents:
            db.add_documents(documents)
            # Mark as indexed in session state
            st.session_state[index_key] = folder_path
        
        return count


def send_email_with_file(
    recipient_email: str,
    recipient_name: str,
    file_path: Path,
    subject: str = None,
    body: str = None
) -> Tuple[bool, str]:
    """
    Send an email with file attachment.
    
    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        # Read all SMTP settings exclusively from env
        sender_name = os.getenv("SMTP_SENDER_NAME", "")
        sender_email = os.getenv("SMTP_SENDER_EMAIL", "")
        sender_password = os.getenv("SMTP_APP_PASSWORD", "")
        smtp_server = os.getenv("SMTP_SERVER", "smtp.gmail.com")
        smtp_port = get_env_int("SMTP_PORT", 465)
        use_starttls = get_env_bool("SMTP_USE_STARTTLS", False)

        if not sender_email:
            return False, "SMTP_SENDER_EMAIL not configured in .env"
        if not sender_password:
            return False, "SMTP_APP_PASSWORD not configured in .env"
        if not recipient_email:
            return False, "Recipient email is required"

        msg = MIMEMultipart()
        msg["From"] = formataddr((sender_name or sender_email, sender_email))
        msg["To"] = formataddr((recipient_name or recipient_email, recipient_email))
        msg["Subject"] = subject or f"Document: {file_path.name}"

        msg.attach(MIMEText(body or os.getenv("SMTP_DEFAULT_BODY", "Hi,\n\nPlease find the attached file.\n"), "plain"))

        # Attach file
        suffix = file_path.suffix.lower()
        with open(file_path, "rb") as f:
            data = f.read()

        if suffix in [".png", ".jpg", ".jpeg", ".gif"]:
            from email.mime.image import MIMEImage
            try:
                image_part = MIMEImage(data, _subtype=suffix.lstrip("."))
            except TypeError:
                image_part = MIMEImage(data)
            image_part.add_header(
                "Content-Disposition",
                "attachment",
                filename=file_path.name,
            )
            msg.attach(image_part)
        else:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(data)
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                "attachment",
                filename=file_path.name,
            )
            msg.attach(part)

        context = ssl.create_default_context()
        port = int(smtp_port)
        # Get timeout from env (default 30 seconds)
        timeout = get_env_int("SMTP_TIMEOUT", 30)
        
        # Set socket default timeout for better network handling
        socket.setdefaulttimeout(timeout)

        if use_starttls:
            with smtplib.SMTP(smtp_server, port, timeout=timeout) as server:
                server.starttls(context=context)
                server.login(sender_email, sender_password)
                server.send_message(msg)
        else:
            with smtplib.SMTP_SSL(smtp_server, port, context=context, timeout=timeout) as server:
                server.login(sender_email, sender_password)
                server.send_message(msg)

        return True, f"Email sent successfully to {recipient_email} with attachment '{file_path.name}'."
    
    except smtplib.SMTPAuthenticationError:
        return False, "Authentication failed. Please verify your email credentials or app password."
    except FileNotFoundError:
        return False, f"File not found: {file_path}"
    except (smtplib.SMTPException, OSError, TimeoutError) as ex:
        error_msg = str(ex)
        if "10060" in error_msg or "timeout" in error_msg.lower() or "timed out" in error_msg.lower():
            return False, f"Connection timeout. Please check your network connection. If using WiFi, try switching to ethernet or check firewall settings. Error: {error_msg}"
        return False, f"Failed to send email: {error_msg}"
    except Exception as ex:
        return False, f"Unexpected error: {str(ex)}"


# Streamlit App
st.set_page_config(
    page_title="Internal Knowledge Navigator",
    page_icon="📚",
    layout="wide"
)

st.title("📚 Internal Knowledge Navigator")

# UI Mode Selector
if "ui_mode" not in st.session_state:
    st.session_state["ui_mode"] = "expert"

mode_col1, mode_col2 = st.columns([1, 10])
with mode_col1:
    ui_mode = st.radio(
        "UI Mode",
        options=["expert", "simplified"],
        index=0 if st.session_state["ui_mode"] == "expert" else 1,
        horizontal=True,
        key="ui_mode_selector"
    )
    st.session_state["ui_mode"] = ui_mode

st.markdown("---")

# User Inputs
# If a folder was selected via the native picker, apply it BEFORE creating the widget
if "pending_folder_path" in st.session_state:
    try:
        st.session_state["folder_path_input"] = st.session_state["pending_folder_path"]
    except Exception:
        pass
    del st.session_state["pending_folder_path"]

# Folder path input (shared between both modes)
folder_cols = st.columns([4, 1])
with folder_cols[0]:
    folder_path = st.text_input(
        "📁 Documents Folder Path",
        placeholder="./documents or C:\\my_docs",
        help="Path to folder containing your PDFs, DOCX, TXT, images",
        key="folder_path_input"
    )
with folder_cols[1]:
    if st.button("Browse…"):
        chosen = browse_for_folder()
        if chosen:
            # Defer setting the widget value until next run, before widget instantiation
            st.session_state["pending_folder_path"] = chosen
            st.rerun()

# Simplified UI Mode
if st.session_state["ui_mode"] == "simplified":
    st.markdown("### 💬 Simplified Mode")
    st.markdown("Ask anything! Examples: 'send mail to user@example.com', 'search for contract.pdf', 'summarize documents'")
    
    # Show folder structure preview in simplified mode
    if folder_path and folder_path.strip():
        with st.expander("📂 Folder Structure Preview", expanded=False):
            structure = get_folder_structure(folder_path, max_depth=2)
            file_counts = get_supported_files_count(folder_path)
            
            if structure:
                if structure[0].startswith("❌"):
                    st.error(structure[0])
                else:
                    # Show file counts
                    if file_counts["total"] > 0:
                        st.success(f"✅ Found {file_counts['total']} supported file(s): "
                                 f"{file_counts['pdf']} PDF(s), "
                                 f"{file_counts['docx']} DOCX(s), "
                                 f"{file_counts['txt']} TXT(s)")
                    else:
                        st.warning("⚠️ No supported files (.pdf, .docx, .txt) found in this folder")
                    
                    # Show folder structure
                    st.markdown("**Folder Structure:**")
                    structure_text = "\n".join(structure)
                    st.code(structure_text, language="text")
            else:
                st.info("Enter a folder path to see its structure")
    
    user_query = st.text_area(
        "❓ Your Query",
        placeholder="e.g., send contract.pdf to avinashdola57@gmail.com, search for contract.pdf, summarize all documents",
        height=120,
        key="simplified_query"
    )
    
    if st.button("🚀 Execute", type="primary", key="simplified_execute"):
        if not folder_path or not folder_path.strip():
            st.error("Please provide the folder path first!")
        elif not user_query or not user_query.strip():
            st.error("Please enter a query!")
        else:
            # Parse query intent
            intent_data = parse_query_intent(user_query)
            intent = intent_data["intent"]
            
            with st.spinner("Processing your request..."):
                if intent == "email":
                    # Handle email intent
                    recipient_email = intent_data.get("email")
                    document_name = intent_data.get("document_name")
                    
                    # For email, we don't need to index - just find the file
                    available_files = list_supported_files(folder_path)
                    
                    if not available_files:
                        st.error("No supported documents found in the folder!")
                    else:
                        
                        # Try to find the specific file mentioned in query
                        selected_file = None
                        if document_name:
                            # Search for the file mentioned in query
                            st.info(f"🔍 Searching for '{document_name}'...")
                            results, suggestions = find_documents(folder_path, document_name)
                            if results:
                                selected_file = results[0]
                                st.success(f"📎 Found file: {selected_file.name}")
                                if len(results) > 1:
                                    st.info(f"💡 Multiple matches found. Sending: {selected_file.name}")
                            else:
                                st.warning(f"⚠️ Could not find exact match for '{document_name}'.")
                                if suggestions:
                                    st.info("💡 Did you mean one of these?")
                                    for idx, (suggested_path, score) in enumerate(suggestions[:3], 1):
                                        similarity_pct = int(score * 100)
                                        st.write(f"{idx}. {suggested_path.name} (Similarity: {similarity_pct}%)")
                                    # Use the most similar suggestion
                                    selected_file = suggestions[0][0]
                                    st.info(f"📎 Using closest match: {selected_file.name}")
                                else:
                                    st.warning("Please select a file below.")
                        
                        if not selected_file:
                            # If no specific file found or not mentioned, use first file or let user select
                            if len(available_files) == 1:
                                selected_file = available_files[0]
                                st.info(f"📎 Using only available file: {selected_file.name}")
                            else:
                                # Show file selector
                                base_path = Path(folder_path)
                                selected_file = st.selectbox(
                                    "📎 Select File to Attach",
                                    options=available_files,
                                    format_func=lambda p: format_file_display(base_path, p),
                                    key="email_file_selector"
                                )
                    
                    if not recipient_email:
                        # Prompt for email if not found
                        st.warning("📧 Email recipient not found in query. Please provide recipient email.")
                        st.info("💡 Tip: Include the email address in your query, e.g., 'send contract.pdf to user@example.com'")
                        
                        recipient_email = st.text_input(
                            "Recipient Email Address", 
                            key="email_recipient_input",
                            placeholder="user@example.com"
                        )
                        
                        if st.button("📧 Send Email", key="send_email_after_prompt"):
                            if recipient_email and selected_file:
                                with st.spinner("Sending email..."):
                                    success, message = send_email_with_file(
                                        recipient_email=recipient_email,
                                        recipient_name="",
                                        file_path=selected_file
                                    )
                                    if success:
                                        st.success(message)
                                    else:
                                        st.error(message)
                            elif not recipient_email:
                                st.error("Please enter a recipient email address.")
                    else:
                        # Email found in query - send immediately
                        if selected_file:
                            with st.spinner("Sending email..."):
                                success, message = send_email_with_file(
                                    recipient_email=recipient_email,
                                    recipient_name="",
                                    file_path=selected_file
                                )
                                if success:
                                    st.success(message)
                                else:
                                    st.error(message)
                        else:
                            st.error("No file selected to send.")
                
                elif intent == "search":
                    # Handle search intent
                    document_name = intent_data.get("document_name")
                    if not document_name:
                        st.error("Please specify which document to search for. Example: 'search for contract.pdf'")
                    else:
                        results, suggestions = find_documents(folder_path, document_name)
                        
                        if results:
                            st.success(f"✅ Found {len(results)} matching document(s):")
                            for idx, file_path in enumerate(results, 1):
                                col1, col2 = st.columns([3, 1])
                                with col1:
                                    file_size = file_path.stat().st_size / 1024
                                    st.markdown(f"**{idx}. {file_path.name}**")
                                    st.caption(f"📍 {file_path.parent}")
                                    st.caption(f"📊 Size: {file_size:.2f} KB")
                                with col2:
                                    if st.button("🚀 Go to", key=f"simplified_goto_{idx}"):
                                        open_file_location(file_path)
                                        st.success(f"Opening location for: {file_path.name}")
                        else:
                            st.warning(f"❌ No documents found matching '{document_name}'")
                            if suggestions:
                                st.markdown("#### 🔍 Did you mean?")
                                for idx, (suggested_path, score) in enumerate(suggestions, 1):
                                    similarity_pct = int(score * 100)
                                    st.markdown(f"**{idx}. {suggested_path.name}** (Similarity: {similarity_pct}%)")
                
                elif intent == "summarize":
                    # Handle summarize intent
                    document_name = intent_data.get("document_name")
                    
                    db = VectorDatabaseManager()
                    # Index documents (will skip if already indexed)
                    docs_indexed = index_folder(folder_path, db)
                    
                    if docs_indexed == 0:
                        st.error("No documents found in the specified folder!")
                    else:
                        rag = RAGEngine()
                        doc_ids_to_summarize = []
                        
                        if document_name:
                            # User specified a document to summarize
                            st.info(f"🔍 Searching for '{document_name}'...")
                            results, suggestions = find_documents(folder_path, document_name)
                            
                            if results:
                                # Found exact match(es)
                                for file_path in results:
                                    # Convert file path to doc_id format
                                    doc_id = str(file_path).replace("\\", "_").replace("/", "_")
                                    # Check if this doc_id exists in the database
                                    all_doc_ids = db.get_all_document_ids()
                                    # Find matching doc_id (might have chunks)
                                    matching_ids = [did for did in all_doc_ids if doc_id in did or did in doc_id]
                                    if matching_ids:
                                        # Use the base doc_id (without chunk suffix)
                                        base_doc_id = doc_id
                                        if base_doc_id not in doc_ids_to_summarize:
                                            doc_ids_to_summarize.append(base_doc_id)
                                
                                if doc_ids_to_summarize:
                                    st.success(f"📎 Found {len(doc_ids_to_summarize)} document(s) to summarize")
                                else:
                                    st.warning(f"⚠️ Found file(s) but couldn't match with database. Trying to find by filename...")
                                    # Try alternative: search by filename in doc_ids
                                    for file_path in results:
                                        filename = file_path.name
                                        all_doc_ids = db.get_all_document_ids()
                                        for doc_id in all_doc_ids:
                                            doc = db.get_document_by_id(doc_id)
                                            if doc and doc.get("metadata"):
                                                if doc["metadata"].get("filename") == filename:
                                                    if doc_id not in doc_ids_to_summarize:
                                                        doc_ids_to_summarize.append(doc_id)
                                                    break
                            elif suggestions:
                                # Use the most similar suggestion
                                file_path = suggestions[0][0]
                                filename = file_path.name
                                st.info(f"📎 Using closest match: {filename}")
                                all_doc_ids = db.get_all_document_ids()
                                for doc_id in all_doc_ids:
                                    doc = db.get_document_by_id(doc_id)
                                    if doc and doc.get("metadata"):
                                        if doc["metadata"].get("filename") == filename:
                                            doc_ids_to_summarize.append(doc_id)
                                            break
                            else:
                                st.error(f"❌ Could not find document matching '{document_name}'")
                                st.info("💡 Available documents in folder:")
                                available_files = list_supported_files(folder_path)
                                for idx, file_path in enumerate(available_files[:10], 1):
                                    st.write(f"{idx}. {file_path.name}")
                        else:
                            # No specific document mentioned - summarize all
                            st.info("🤖 Generating summary of all documents...")
                            all_doc_ids = db.get_all_document_ids()
                            doc_ids_to_summarize = all_doc_ids
                        
                        if doc_ids_to_summarize:
                            st.info(f"📄 Analyzing {len(doc_ids_to_summarize)} document(s)...")
                            # Summarize specified documents
                            result = rag.summarize(doc_ids_to_summarize, summary_type="abstractive")
                            
                            if result.get("status") == "success":
                                summary_title = f"Summary of {len(doc_ids_to_summarize)} Document(s)" if document_name else "Summary of All Documents"
                                st.success("✅ Summary Complete!")
                                st.markdown(f"### 📊 {summary_title}")
                                st.write(result.get("summary", "No summary available"))
                            else:
                                st.error(f"Error generating summary: {result.get('message', 'Unknown error')}")
                        else:
                            st.error("No documents found to summarize.")
                
                else:
                    # Default: RAG query
                    db = VectorDatabaseManager()
                    # Index documents (will skip if already indexed)
                    docs_indexed = index_folder(folder_path, db)
                    
                    if docs_indexed == 0:
                        st.error("No documents found in the specified folder!")
                    else:
                        rag = RAGEngine()
                        doc_ids_to_query = None
                        
                        # Check if query mentions a specific document (e.g., "explain project1.pdf" or "explain project1")
                        # Try to extract document name from query
                        query_lower = user_query.lower()
                        
                        # Pattern to match filenames (with extensions)
                        filename_pattern = r'\b[\w\-_]+\.(pdf|docx|txt|doc)\b'
                        filenames = re.findall(filename_pattern, user_query, re.IGNORECASE)
                        
                        # Also try to extract potential document names without extensions
                        # Look for words that might be document names (after common verbs like explain, what, tell, etc.)
                        doc_name = None
                        if filenames:
                            # Found a filename with extension
                            filename_match = re.search(r'\b([\w\-_]+\.(?:pdf|docx|txt|doc))\b', user_query, re.IGNORECASE)
                            if filename_match:
                                doc_name = filename_match.group(1)
                        else:
                            # Try to extract potential document name (word after common query verbs)
                            # Remove common query words
                            query_words = re.sub(r'\b(explain|what|is|tell|me|about|describe|show|the|a|an)\b', '', query_lower, flags=re.IGNORECASE)
                            query_words = query_words.strip()
                            # Take the first significant word/phrase (up to 2 words)
                            words = query_words.split()
                            if words:
                                # Take first 1-2 words as potential document name
                                potential_name = ' '.join(words[:2]).strip(' ,.')
                                if len(potential_name) > 2:  # Only if it's a meaningful name
                                    doc_name = potential_name
                        
                        if doc_name:
                            st.info(f"🔍 Searching for document: {doc_name}")
                            
                            # Search for the document in the folder (with fuzzy matching)
                            results, suggestions = find_documents(folder_path, doc_name)
                            
                            # Use suggestions if no exact match found (handles spelling errors)
                            if not results and suggestions:
                                # Use the most similar suggestion (fuzzy match)
                                file_path = suggestions[0][0]
                                filename_only = file_path.name
                                similarity_score = suggestions[0][1]
                                if similarity_score >= 0.5:  # Only use if reasonably similar
                                    st.info(f"📎 Found closest match: {filename_only} (similarity: {int(similarity_score * 100)}%)")
                                    results = [file_path]
                            
                            if results:
                                # Found the document, get its doc_id
                                doc_ids_to_query = []
                                for file_path in results:
                                    doc_id = str(file_path).replace("\\", "_").replace("/", "_")
                                    # Check if this doc_id exists in the database
                                    all_doc_ids = db.get_all_document_ids()
                                    # Find matching doc_id
                                    matching_ids = [did for did in all_doc_ids if doc_id == did or did.startswith(doc_id + "_chunk")]
                                    if matching_ids:
                                        # Use the base doc_id (without chunk suffix)
                                        base_doc_id = doc_id
                                        if base_doc_id not in doc_ids_to_query:
                                            doc_ids_to_query.append(base_doc_id)
                                
                                if not doc_ids_to_query:
                                    # Try alternative: search by filename in doc_ids
                                    for file_path in results:
                                        filename_only = file_path.name
                                        all_doc_ids = db.get_all_document_ids()
                                        for doc_id in all_doc_ids:
                                            doc = db.get_document_by_id(doc_id)
                                            if doc and doc.get("metadata"):
                                                if doc["metadata"].get("filename") == filename_only:
                                                    if doc_id not in doc_ids_to_query:
                                                        doc_ids_to_query.append(doc_id)
                                                    break
                                
                                if doc_ids_to_query:
                                    st.success(f"📎 Using document: {results[0].name}")
                        
                        st.info("🤖 Analyzing documents with AI...")
                        result = rag.query(user_query, document_ids=doc_ids_to_query)
                        
                        st.success("✅ Analysis Complete!")
                        st.markdown("### 📊 Results")
                        st.markdown("#### Summary")
                        st.write(result.get("summary", "No summary available"))
    
    st.markdown("---")
    st.info("💡 **Tips:**\n- Say 'send mail to email@example.com' to email documents\n- Say 'search for filename' to find documents\n- Say 'summarize' to get document summaries\n- Ask any question about your documents!")

# Expert UI Mode
else:
    st.markdown("### ⚙️ Expert Mode")
    
    # Show folder structure preview
    if folder_path and folder_path.strip():
        with st.expander("📂 Folder Structure Preview", expanded=True):
            structure = get_folder_structure(folder_path, max_depth=2)
            file_counts = get_supported_files_count(folder_path)
            
            if structure:
                if structure[0].startswith("❌"):
                    st.error(structure[0])
                else:
                    # Show file counts
                    if file_counts["total"] > 0:
                        st.success(f"✅ Found {file_counts['total']} supported file(s): "
                                 f"{file_counts['pdf']} PDF(s), "
                                 f"{file_counts['docx']} DOCX(s), "
                                 f"{file_counts['txt']} TXT(s)")
                    else:
                        st.warning("⚠️ No supported files (.pdf, .docx, .txt) found in this folder")
                    
                    # Show folder structure
                    st.markdown("**Folder Structure:**")
                    structure_text = "\n".join(structure)
                    st.code(structure_text, language="text")
            else:
                st.info("Enter a folder path to see its structure")

    # Find Document Feature
    st.markdown("---")
    st.markdown("### 🔍 Find Document")
    st.markdown("Search for a specific document in the folder and open its location")

    document_name = st.text_input(
        "📄 Document Name",
        placeholder="e.g., contract_v1.pdf or project",
        help="Enter the document name (partial match, case-insensitive)",
        key="document_name"
    )

    if st.button("🔎 Find Document", key="find_doc_button"):
        if not folder_path or not folder_path.strip():
            st.error("Please provide the folder path above before searching.")
        elif not document_name or not document_name.strip():
            st.error("Please provide a document name to search for.")
        else:
            with st.spinner("Searching for documents..."):
                results, suggestions = find_documents(folder_path, document_name)
                
                if results:
                    st.success(f"✅ Found {len(results)} matching document(s):")
                    
                    for idx, file_path in enumerate(results, 1):
                        col1, col2, col3 = st.columns([3, 1, 1])
                        
                        with col1:
                            # Show file info
                            file_size = file_path.stat().st_size / 1024  # Size in KB
                            st.markdown(f"**{idx}. {file_path.name}**")
                            st.caption(f"📍 {file_path.parent}")
                            st.caption(f"📊 Size: {file_size:.2f} KB | Type: {file_path.suffix.upper()}")
                        
                        with col2:
                            # Show full path in expander
                            with st.expander("📋 Full Path"):
                                st.code(str(file_path), language="text")
                        
                        with col3:
                            # Go to button
                            if st.button("🚀 Go to", key=f"goto_{idx}_{file_path.name}"):
                                open_file_location(file_path)
                                st.success(f"Opening location for: {file_path.name}")
                    
                    st.markdown("---")
                else:
                    st.warning(f"❌ No documents found matching '{document_name}' in the specified folder.")
                    if suggestions:
                        st.markdown("#### 🔍 Did you mean?")
                        for idx, (suggested_path, score) in enumerate(suggestions, 1):
                            similarity_pct = int(score * 100)
                            col1, col2 = st.columns([3, 1])

                            with col1:
                                st.markdown(f"**{idx}. {suggested_path.name}**  ")
                                st.caption(f"📍 {suggested_path.parent}")
                                st.caption(f"🔎 Similarity: {similarity_pct}%")

        st.info("💡 Tip: Try refining the document name or use one of the suggested files above.")
    else:
        st.info("💡 Tip: Try a partial name match. The search is case-insensitive.")
    st.markdown("---")

    if not folder_path or not folder_path.strip():
        st.info("Provide the folder path above to load documents for emailing.")
    else:
        available_files = list_supported_files(folder_path)

        if not available_files:
            st.warning("No supported documents (.pdf, .docx, .txt) found to send.")
        else:
            base_path = Path(folder_path)
            with st.form("email_sender_form"):
                selected_file = st.selectbox(
                    "📎 File to Attach",
                    options=available_files,
                    format_func=lambda p: format_file_display(base_path, p),
                )

                recipient_name = st.text_input("Recipient Name", value="")
                recipient_email = st.text_input("Recipient Email Address")

                # Subject defaults to filename
                default_subject = f"Document: {selected_file.name}"
                email_subject = st.text_input("Email Subject", value=default_subject)
                email_body_default = os.getenv(
                    "SMTP_DEFAULT_BODY", "Hi,\n\nPlease find the attached file.\n"
                )
                email_body = st.text_area(
                    "Email Message",
                    value=email_body_default,
                    height=120,
                )

                send_email = st.form_submit_button("Send Email ✉️")

                if send_email:
                    # Read all SMTP settings exclusively from env
                    sender_name = os.getenv("SMTP_SENDER_NAME", "")
                    sender_email = os.getenv("SMTP_SENDER_EMAIL", "")
                    sender_password = os.getenv("SMTP_APP_PASSWORD", "")
                    smtp_server = os.getenv("SMTP_SERVER", "smtp.gmail.com")
                    smtp_port = get_env_int("SMTP_PORT", 465)
                    use_starttls = get_env_bool("SMTP_USE_STARTTLS", False)

                    missing_fields = []
                    if not sender_email:
                        missing_fields.append("SMTP_SENDER_EMAIL (.env)")
                    if not sender_password:
                        missing_fields.append("SMTP_APP_PASSWORD (.env)")
                    if not recipient_email:
                        missing_fields.append("Recipient Email")

                    if missing_fields:
                        st.error("Missing required settings: " + ", ".join(missing_fields))
                    else:
                        try:
                            msg = MIMEMultipart()
                            msg["From"] = formataddr((sender_name or sender_email, sender_email))
                            msg["To"] = formataddr((recipient_name or recipient_email, recipient_email))
                            msg["Subject"] = email_subject or default_subject

                            msg.attach(MIMEText(email_body or "", "plain"))

                            # Attach file (image types handled specially)
                            suffix = selected_file.suffix.lower()
                            with open(selected_file, "rb") as f:
                                data = f.read()

                            if suffix in [".png", ".jpg", ".jpeg", ".gif"]:
                                from email.mime.image import MIMEImage

                                try:
                                    image_part = MIMEImage(data, _subtype=suffix.lstrip("."))
                                except TypeError:
                                    image_part = MIMEImage(data)
                                image_part.add_header(
                                    "Content-Disposition",
                                    "attachment",
                                    filename=selected_file.name,
                                )
                                msg.attach(image_part)
                            else:
                                part = MIMEBase("application", "octet-stream")
                                part.set_payload(data)
                                encoders.encode_base64(part)
                                part.add_header(
                                    "Content-Disposition",
                                    "attachment",
                                    filename=selected_file.name,
                                )
                                msg.attach(part)

                            context = ssl.create_default_context()
                            port = int(smtp_port)
                            # Get timeout from env (default 30 seconds)
                            timeout = get_env_int("SMTP_TIMEOUT", 30)
                            
                            # Set socket default timeout for better network handling
                            socket.setdefaulttimeout(timeout)

                            if use_starttls:
                                with smtplib.SMTP(smtp_server, port, timeout=timeout) as server:
                                    server.starttls(context=context)
                                    server.login(sender_email, sender_password)
                                    server.send_message(msg)
                            else:
                                with smtplib.SMTP_SSL(smtp_server, port, context=context, timeout=timeout) as server:
                                    server.login(sender_email, sender_password)
                                    server.send_message(msg)

                        except smtplib.SMTPAuthenticationError:
                            st.error(
                                "Authentication failed. Please verify your email credentials or app password."
                            )
                        except FileNotFoundError:
                            st.error("Selected file could not be read. Please verify it still exists.")
                        except (smtplib.SMTPException, OSError, TimeoutError) as ex:
                            error_msg = str(ex)
                            if "10060" in error_msg or "timeout" in error_msg.lower() or "timed out" in error_msg.lower():
                                st.error(f"Connection timeout. Please check your network connection. If using WiFi, try switching to ethernet or check firewall settings. Error: {error_msg}")
                            else:
                                st.error(f"Failed to send email: {error_msg}")
                        except Exception as ex:
                            st.error(f"Unexpected error while sending email: {str(ex)}")
                        else:
                            st.success(
                                f"Email sent successfully to {recipient_email} with attachment '{selected_file.name}'."
                            )

    st.markdown("---")

    query = st.text_area(
        "❓ Your Query",
        placeholder="e.g., Compare risk clauses between contract v1 and v2",
        height=100
    )

    if st.button("🔍 Search & Analyze", type="primary"):
        if not folder_path or not query:
            st.error("Please provide both folder path and query!")
        else:
            with st.spinner("Processing..."):
                # Check if query is a summarize request
                intent_data = parse_query_intent(query)
                is_summarize = intent_data["intent"] == "summarize"
                
                # Save input.json
                input_data = {
                    "folder_path": folder_path,
                    "query": query,
                    "timestamp": datetime.now().isoformat(),
                    "intent": intent_data["intent"]
                }
                with open("input.json", "w") as f:
                    json.dump(input_data, f, indent=2)
                
                # Index documents (will skip if already indexed)
                db = VectorDatabaseManager()
                docs_indexed = index_folder(folder_path, db)
                
                if docs_indexed == 0:
                    st.error("No documents found in the specified folder!")
                    st.stop()
                
                # Update metadata.json
                metadata = {
                    "project": "Knowledge Navigator",
                    "last_run": datetime.now().isoformat(),
                    "folder_indexed": folder_path,
                    "total_documents": docs_indexed,
                    "file_types": [".pdf", ".docx", ".txt"]
                }
                with open("metadata.json", "w") as f:
                    json.dump(metadata, f, indent=2)
                
                if is_summarize:
                    # Handle summarize intent
                    document_name = intent_data.get("document_name")
                    rag = RAGEngine()
                    doc_ids_to_summarize = []
                    
                    if document_name:
                        # User specified a document to summarize
                        st.info(f"🔍 Searching for '{document_name}'...")
                        results, suggestions = find_documents(folder_path, document_name)
                        
                        if results:
                            # Found exact match(es)
                            for file_path in results:
                                # Convert file path to doc_id format
                                doc_id = str(file_path).replace("\\", "_").replace("/", "_")
                                # Check if this doc_id exists in the database
                                all_doc_ids = db.get_all_document_ids()
                                # Find matching doc_id (might have chunks)
                                matching_ids = [did for did in all_doc_ids if doc_id in did or did in doc_id]
                                if matching_ids:
                                    # Use the base doc_id (without chunk suffix)
                                    base_doc_id = doc_id
                                    if base_doc_id not in doc_ids_to_summarize:
                                        doc_ids_to_summarize.append(base_doc_id)
                            
                            if doc_ids_to_summarize:
                                st.success(f"📎 Found {len(doc_ids_to_summarize)} document(s) to summarize")
                            else:
                                st.warning(f"⚠️ Found file(s) but couldn't match with database. Trying to find by filename...")
                                # Try alternative: search by filename in doc_ids
                                for file_path in results:
                                    filename = file_path.name
                                    all_doc_ids = db.get_all_document_ids()
                                    for doc_id in all_doc_ids:
                                        doc = db.get_document_by_id(doc_id)
                                        if doc and doc.get("metadata"):
                                            if doc["metadata"].get("filename") == filename:
                                                if doc_id not in doc_ids_to_summarize:
                                                    doc_ids_to_summarize.append(doc_id)
                                                break
                        elif suggestions:
                            # Use the most similar suggestion
                            file_path = suggestions[0][0]
                            filename = file_path.name
                            st.info(f"📎 Using closest match: {filename}")
                            all_doc_ids = db.get_all_document_ids()
                            for doc_id in all_doc_ids:
                                doc = db.get_document_by_id(doc_id)
                                if doc and doc.get("metadata"):
                                    if doc["metadata"].get("filename") == filename:
                                        doc_ids_to_summarize.append(doc_id)
                                        break
                        else:
                            st.error(f"❌ Could not find document matching '{document_name}'")
                            st.info("💡 Available documents in folder:")
                            available_files = list_supported_files(folder_path)
                            for idx, file_path in enumerate(available_files[:10], 1):
                                st.write(f"{idx}. {file_path.name}")
                    else:
                        # No specific document mentioned - summarize all
                        st.info("🤖 Generating summary of all documents...")
                        all_doc_ids = db.get_all_document_ids()
                        doc_ids_to_summarize = all_doc_ids
                    
                    if not doc_ids_to_summarize:
                        st.error("No documents found in the database!")
                        st.stop()
                    
                    st.info(f"📄 Analyzing {len(doc_ids_to_summarize)} document(s)...")
                    # Summarize specified documents
                    result = rag.summarize(doc_ids_to_summarize, summary_type="abstractive")
                    
                    if result.get("status") == "success":
                        summary_title = f"Summary of {len(doc_ids_to_summarize)} Document(s)" if document_name else "Summary of All Documents"
                        
                        # Save output.json
                        output_data = {
                            "status": "success",
                            "query": query,
                            "intent": "summarize",
                            "documents_analyzed": result.get("documents_analyzed", 0),
                            "results": result
                        }
                        with open("output.json", "w") as f:
                            json.dump(output_data, f, indent=2)
                        
                        # Display results
                        st.success("✅ Summary Complete!")
                        st.markdown(f"### 📊 {summary_title}")
                        st.write(result.get("summary", "No summary available"))
                        
                        # Show full JSON
                        with st.expander("📄 View Full Results (JSON)"):
                            st.json(output_data)
                        
                        # Download button
                        st.download_button(
                            "💾 Download Results (JSON)",
                            data=json.dumps(output_data, indent=2),
                            file_name=f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                            mime="application/json"
                        )
                    else:
                        st.error(f"Error generating summary: {result.get('message', 'Unknown error')}")
                else:
                    # Execute RAG query (default behavior)
                    rag = RAGEngine()
                    doc_ids_to_query = None
                    
                    # Check if query mentions a specific document (e.g., "explain project1.pdf" or "explain project1")
                    query_lower = query.lower()
                    
                    # Pattern to match filenames (with extensions)
                    filename_pattern = r'\b[\w\-_]+\.(pdf|docx|txt|doc)\b'
                    filenames = re.findall(filename_pattern, query, re.IGNORECASE)
                    
                    # Also try to extract potential document names without extensions
                    doc_name = None
                    if filenames:
                        # Found a filename with extension
                        filename_match = re.search(r'\b([\w\-_]+\.(?:pdf|docx|txt|doc))\b', query, re.IGNORECASE)
                        if filename_match:
                            doc_name = filename_match.group(1)
                    else:
                        # Try to extract potential document name (word after common query verbs)
                        # Remove common query words
                        query_words = re.sub(r'\b(explain|what|is|tell|me|about|describe|show|the|a|an)\b', '', query_lower, flags=re.IGNORECASE)
                        query_words = query_words.strip()
                        # Take the first significant word/phrase (up to 2 words)
                        words = query_words.split()
                        if words:
                            # Take first 1-2 words as potential document name
                            potential_name = ' '.join(words[:2]).strip(' ,.')
                            if len(potential_name) > 2:  # Only if it's a meaningful name
                                doc_name = potential_name
                    
                    if doc_name:
                        st.info(f"🔍 Searching for document: {doc_name}")
                        
                        # Search for the document in the folder (with fuzzy matching)
                        results, suggestions = find_documents(folder_path, doc_name)
                        
                        # Use suggestions if no exact match found (handles spelling errors)
                        if not results and suggestions:
                            # Use the most similar suggestion (fuzzy match)
                            file_path = suggestions[0][0]
                            filename_only = file_path.name
                            similarity_score = suggestions[0][1]
                            if similarity_score >= 0.5:  # Only use if reasonably similar
                                st.info(f"📎 Found closest match: {filename_only} (similarity: {int(similarity_score * 100)}%)")
                                results = [file_path]
                        
                        if results:
                            # Found the document, get its doc_id
                            doc_ids_to_query = []
                            for file_path in results:
                                doc_id = str(file_path).replace("\\", "_").replace("/", "_")
                                # Check if this doc_id exists in the database
                                all_doc_ids = db.get_all_document_ids()
                                # Find matching doc_id
                                matching_ids = [did for did in all_doc_ids if doc_id == did or did.startswith(doc_id + "_chunk")]
                                if matching_ids:
                                    # Use the base doc_id (without chunk suffix)
                                    base_doc_id = doc_id
                                    if base_doc_id not in doc_ids_to_query:
                                        doc_ids_to_query.append(base_doc_id)
                            
                            if not doc_ids_to_query:
                                # Try alternative: search by filename in doc_ids
                                for file_path in results:
                                    filename_only = file_path.name
                                    all_doc_ids = db.get_all_document_ids()
                                    for doc_id in all_doc_ids:
                                        doc = db.get_document_by_id(doc_id)
                                        if doc and doc.get("metadata"):
                                            if doc["metadata"].get("filename") == filename_only:
                                                if doc_id not in doc_ids_to_query:
                                                    doc_ids_to_query.append(doc_id)
                                                break
                            
                            if doc_ids_to_query:
                                st.success(f"📎 Using document: {results[0].name}")
                    
                    st.info("🤖 Analyzing documents with AI...")
                    result = rag.query(query, document_ids=doc_ids_to_query)
                    
                    # Save output.json
                    output_data = {
                        "status": "success",
                        "query": query,
                        "execution_time": result.get("execution_time", 0),
                        "results": result
                    }
                    with open("output.json", "w") as f:
                        json.dump(output_data, f, indent=2)
                    
                    # Display results
                    st.success("✅ Analysis Complete!")
                    st.markdown("### 📊 Results")
                    
                    # Show summary
                    st.markdown("#### Summary")
                    st.write(result.get("summary", "No summary available"))
                    
                    
                    # Show full JSON
                    with st.expander("📄 View Full Results (JSON)"):
                        st.json(output_data)
                    
                    # Download button
                    st.download_button(
                        "💾 Download Results (JSON)",
                        data=json.dumps(output_data, indent=2),
                        file_name=f"results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json"
                    )

# Sidebar with info
with st.sidebar:
    st.header("ℹ️ About")
    st.markdown("""
    This tool helps you:
    - 🔍 Search documents semantically
    - 📊 Compare document versions
    - 📝 Extract key information
    - 🤖 Get AI-powered insights
    - 🚀 Find and locate documents quickly
    """)
    
    st.header("📋 Supported Formats")
    st.markdown("""
    - ✅ PDF (.pdf)
    - ✅ Word Documents (.docx)
    - ✅ Text Files (.txt)
    """)
    
    # Show metadata if exists
    if os.path.exists("metadata.json"):
        st.header("📊 Last Run Info")
        try:
            with open("metadata.json", "r") as f:
                meta = json.load(f)
                st.json(meta)
        except:
            pass
